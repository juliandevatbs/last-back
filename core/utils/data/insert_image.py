from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Pt, Inches


def insert_image (search_text: str, doc, image_path: str, width_inches = 3.0, height_inches = 2.0, alignment = 'center', space_before=Pt(10),  space_after=Pt(10)) -> bool:




    """

    INSERTS A AFTER A SPECIFIC PARAGRAPH

    :param search_text:
    :param image_path:
    :param width_inches:
    :param height_inches:
    :param alignment:
    :param space_before:
    :param space_after:
    :return: bool True if correct
    """



    try:


        target_paragraph = None

        for paragraph in doc.paragraphs:

            if search_text in paragraph.text:

                target_paragraph = paragraph

                break

        if not target_paragraph:

            print(f"PARAGRAPH NOT FOUND")
            return False



        # Verify if the image exists
        new_para = doc.add_paragraph()


        new_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        new_para.paragraph_format.space_before = space_before
        new_para.paragraph_format.space_after = space_after

        run = new_para.add_run()
        run.add_picture(image_path, width = Inches(width_inches), height=Inches(height_inches))

        target_element = target_paragraph._element
        parent = target_element.getparent()
        parent.insert(parent.index(target_element) + 1, new_para.element)


        return True


    except Exception as e:

        return False




