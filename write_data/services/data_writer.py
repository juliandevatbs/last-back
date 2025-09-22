import copy
import os
import subprocess
from importlib.metadata import PackageNotFoundError
from sys import platform

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.shared import Pt, Inches
from docx.text.paragraph import Paragraph
from docx.oxml.ns import qn


# Import custom exceptions
from core.exceptions import NoDataError
from core.services.server_service import ServerService
from core.utils.data.generate_objective_paragraph import generate_objective_paragraph
from core.utils.data.get_today_date import get_today_date
from core.utils.data.split_department_municipality import split_department_municipality
from core.utils.data.int_to_string_relative import int_to_string_relative
from write_data.services.text_fixing import sampling_site_fixing
from datetime import datetime


# Service for write template
class WordService:


    def __init__(self, template_to_write, data_to_write: dict):

        self.doc = None
        self.data_to_write = data_to_write
        self.template_to_write = template_to_write

        # Server instance
        self.server_templates = ServerService()

        #IMPORTANT CONSTRAINTS
        self.font = "Century Gothic"



        # Get data from the dict
        self.main_data = self.data_to_write["main_data"]
        self.sampling_data = self.data_to_write["sampling_data"]
        self.samples = self.data_to_write["samples"]
        self.surveillance_data = self.data_to_write["surveillance_data"]


        # General data for pages
        self.sampling_site = self.sampling_data["sampling_site"]
        self.prefix_water_title = "CARACTERIZACIÓN FISICOQUÍMICA DE"
        self.type_prefix = "FISICOQUÍMICA"
        self.water_type = self.sampling_data["sampling_site"]
        self.report_name = f"{self.prefix_water_title} {self.water_type}"
        self.fixed_sampling_site = sampling_site_fixing(self.sampling_site)
        self.client_contact_name = self.main_data["contact_client_name"]
        self.client_name = self.main_data["client_name"]
        self.report_number = self.main_data["report_number"]
        self.sampling_date = self.sampling_data["sampling_date"]
        self.sampling_date_fixed = self.sampling_date.strftime("%Y-%m-%d")
        self.samples_quantity = len(self.samples)
        self.municipality = self.main_data["municipality"]
        self.simple_water_type = self.surveillance_data["water_type"]


        # License data
        self.license_number = "Nº160 PZ-RES1712-6998"



    #Function to clear cell

    def clear_cell(self, cell):

        """
            Limpia la celda pero preserva las imágenes
            """
        # Guardar referencias a las imágenes
        images = []
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                if run._element.xpath('.//w:drawing'):  # Tiene imagen
                    images.extend(run._element.xpath('.//w:drawing'))

        # Limpiar párrafos normalmente
        for p in cell.paragraphs:
            p._element.getparent().remove(p._element)

        # Si había imágenes, crear un párrafo y restaurarlas
        if images:
            p = cell.add_paragraph()
            for img in images:
                # Agregar la imagen de vuelta
                p._element.append(img)


    #Function to write cell with styles and constraints
    def write_cell(self, cell, text, size=10, bold=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after = 0):

        """
            Text can be
                string simple: "Normal text"
                tuple list for mixed format : [("Bold text": True), ("Normal text": False)]

        :param cell:
        :param text:
        :param size:
        :param bold:
        :param align:
        :param space_after:
        :return:

        """




        self.clear_cell(cell)

        p= cell.add_paragraph()


        if isinstance(text, list):

            #Mixed format - each element is a tuple (text, bold)
            for text_part, is_bold in text:

                run = p.add_run(text_part)
                run.font.name = self.font
                run.font.size = Pt(size)
                run.bold = is_bold
        else:

            run = p.add_run(text)
            run.font.name = self.font
            run.font.size = Pt(size)
            run.bold = bold

        #Format
        p.alignment = align


        p.paragraph_format.space_before= Pt(0)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1









    def validate_template(self):

        # Charge and validate the template. If is not valid launch a exception

        try:



            # Open the docx
            self.doc = self.template_to_write


            _ = self.doc.paragraphs

            print("Valid template")
            return True

        except (PackageNotFoundError, Exception) as ex:


            print(f"Error while validating the template {ex}")
            self.doc = None
            return False

    def write_specific_objectives(self):

        return True

    def write_goals(self):


        relative_quantity_samples = int_to_string_relative(self.samples_quantity)

        # Number of samples plural o singular
        number_samples_p_s = 'punto'

        if self.samples_quantity > 1:

            number_samples_p_s = "puntos"

        return True

    def write_first_table(self):

        if not self.template_to_write.tables:
            print("No hay tablas")
            return

        first_table = self.template_to_write.tables[0]

        titles = [

            "INFORME DE MONITOREO",
            f"{self.prefix_water_title} {self.fixed_sampling_site}",
            f"{self.fixed_sampling_site}"


        ]

        combined_text = "\n\n\n".join(titles)

        self.write_cell(first_table.cell(0,0), combined_text, size=20, bold=True, space_after=30)


        # ==== CREAR SUBTABLA NUEVA DESPUÉS DE LOS 3 TITULOS ====
        # agregamos una fila nueva
        first_table.add_row()
        new_cell = first_table.cell(1, 0)

        # crear una subtabla en esa celda
        subtable = new_cell.add_table(rows=1, cols=3)  # ejemplo: 3 columnas
        subtable.columns[0].width = Inches(2.5)
        subtable.columns[1].width = Inches(2.5)
        subtable.columns[2].width = Inches(2.5)
        subtable.style = "Table Grid"


        for row in subtable.rows:

            for i, width in enumerate([Inches(2.5), Inches(2.5), Inches(2.5)]):

                row.cells[i].width = width

        # llenar subtabla con datos

        self.write_cell(subtable.cell(0, 0), [("Presentado a: ", True), (f"\n{self.client_name}", False), (f"\n{self.client_contact_name}", False)], size=11)
        self.write_cell(subtable.cell(0, 1), [("Revisador por: ", True), (f"\nJulian Criollo", False)], size=11)
        self.write_cell(subtable.cell(0, 2), [("Autorizado por: ", True), ("\nJulian Criollo", False)], size=11)


        # After the table - Report data
        # Report number
        self.write_cell(
            first_table.cell(2, 0),
        [
                ("INFORME NUMERO: ", True),
                ("1702", False)

            ],
                size=11,
                bold=False
        )


        self.write_cell(first_table.cell(4, 0),
                        [
                            ("\n\n\n\n\n\n\n\nFecha de emisión de informe: ", True),
                            (get_today_date(), False)
                        ],
                        size=11,
                        bold=False


                        )




    def write_objective(self):

        print(self.sampling_site)
        print(self.client_name)
        print(self.type_prefix)
        print(self.water_type)
        print(self.license_number)

        objective_paragraph = generate_objective_paragraph(
            self.fixed_sampling_site.lower(),
            self.client_name,
            self.type_prefix.lower(),
            self.simple_water_type.lower(),
            self.license_number,
            None,
            split_department_municipality(self.municipality, part="first"),
            split_department_municipality(self.municipality, part="second"),
            None
        )

        for i, paragraph in enumerate(self.doc.paragraphs):
            if "OBJETIVO GENERAL" in paragraph.text.upper():
                print(f"Encontrado en párrafo {i}: '{paragraph.text}'")

                # Verificar si hay un párrafo siguiente
                if i + 1 < len(self.doc.paragraphs):
                    # Insertar ANTES del siguiente párrafo
                    next_paragraph = self.doc.paragraphs[i + 1]
                    new_paragraph = next_paragraph.insert_paragraph_before(objective_paragraph)
                else:
                    # Si es el último párrafo, agregar al final del documento
                    new_paragraph = self.doc.add_paragraph(objective_paragraph)

                # Formato
                new_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                new_paragraph.paragraph_format.space_before = Pt(6)
                new_paragraph.paragraph_format.space_after = Pt(6)

                for run in new_paragraph.runs:
                    run.font.name = self.font
                    run.font.size = Pt(11)
                    run.bold = False

                print("✅ Objetivo escrito correctamente")
                return True

        print("❌ No se encontró 'OBJETIVO GENERAL'")
        return False

    def insert_sampling_methodology_text(self):
            """
            Inserta texto de metodología de muestreo DESPUÉS del título "2.4.1. Metodología de muestreo"
            """
            # Texto con las variables dinámicas
            methodology_text = f"""En el estudio de calidad fisicoquímica de agua superficial efectuada en los puntos de monitoreo establecidos en el área de influencia de {self.fixed_sampling_site} de {self.client_name}, se tomaron las muestras a través de la recolección directa a una hora determinada, reflejando así las características físicas y químicas instantáneas de los cuerpos de agua; para esto se siguió el Procedimiento de Toma de Muestras de Aguas establecido en el Laboratorio ChemiLab (PGC 04 004), el cual se encuentra acreditado según Resolución 1042 de 26 de septiembre de 2024 del Instituto de Hidrología, Meteorología y Estudios Ambientales (IDEAM), por la cual se dictan unas disposiciones con respecto a la vigencia de la acreditación para los laboratorios ambientales."""

            # Buscar título "2.4.1. Metodología de muestreo" o variantes
            search_patterns = [
                "2.4.1. METODOLOGÍA DE MUESTREO",
                "2.4.1 METODOLOGÍA DE MUESTREO",
                "2.4.1.METODOLOGÍA DE MUESTREO",
                "METODOLOGÍA DE MUESTREO",
                "2.4.1. Metodología de muestreo",
                "2.4.1 Metodología de muestreo"
            ]

            for i, paragraph in enumerate(self.doc.paragraphs):
                paragraph_text = paragraph.text.strip().upper()

                if any(pattern in paragraph_text for pattern in search_patterns):
                    print(f"✅ Encontrado título de metodología: '{paragraph.text}' en párrafo {i}")

                    # Insertar después del título
                    if i + 1 < len(self.doc.paragraphs):
                        next_paragraph = self.doc.paragraphs[i + 1]
                        new_paragraph = next_paragraph.insert_paragraph_before(methodology_text)
                    else:
                        new_paragraph = self.doc.add_paragraph(methodology_text)

                    # Formatear el párrafo
                    new_paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    new_paragraph.paragraph_format.space_before = Pt(6)
                    new_paragraph.paragraph_format.space_after = Pt(6)
                    new_paragraph.paragraph_format.first_line_indent = Pt(18)  # Sangría primera línea

                    # Aplicar fuente y tamaño
                    for run in new_paragraph.runs:
                        run.font.name = self.font
                        run.font.size = Pt(11)
                        run.bold = False

                    print("✅ Texto de metodología de muestreo insertado correctamente")
                    return True

            print("❌ No se encontró el título '2.4.1. Metodología de muestreo'")
            return False
        #self.explore_tables(first_table)

    def recreate_table_of_contents(self):
        """
        Recrea completamente la tabla de contenido
        """
        try:
            # 1. Primero encontrar y eliminar el TOC existente
            toc_found = False
            elements_to_remove = []

            for element in self.doc.element.body:
                if element.tag.endswith('}p'):
                    para_text = ""
                    for t in element.findall('.//w:t',
                                             {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                        if t.text:
                            para_text += t.text

                    if 'TABLA DE CONTENIDO' in para_text.upper() or 'CONTENIDO' in para_text.upper():
                        toc_found = True
                        # Marcar este elemento y los siguientes para eliminar
                        elements_to_remove.append(element)

                elif toc_found and element.tag.endswith('}p'):
                    # Seguir buscando párrafos hasta encontrar el fin del TOC
                    elements_to_remove.append(element)
                    para_text = ""
                    for t in element.findall('.//w:t',
                                             {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                        if t.text:
                            para_text += t.text

                    # Si encontramos el final del TOC (normalmente una página en blanco o nuevo título)
                    if not para_text.strip() or 'INTRODUCCIÓN' in para_text.upper() or '1.' in para_text:
                        break

            # 2. Eliminar el TOC antiguo
            for element in elements_to_remove:
                element.getparent().remove(element)

            print(f"✅ TOC antiguo eliminado ({len(elements_to_remove)} elementos)")

            # 3. Buscar donde insertar el nuevo TOC (después del título)
            insert_position = None
            for i, element in enumerate(self.doc.element.body):
                if element.tag.endswith('}p'):
                    para_text = ""
                    for t in element.findall('.//w:t',
                                             {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                        if t.text:
                            para_text += t.text

                    if 'TABLA DE CONTENIDO' in para_text.upper() or 'CONTENIDO' in para_text.upper():
                        insert_position = i + 1
                        break

            # 4. Crear nuevo campo TOC
            if insert_position is not None:
                self._create_toc_field(insert_position)
                print("✅ Nuevo campo TOC creado")
                return True

            return False

        except Exception as e:
            print(f"❌ Error al recrear tabla de contenido: {e}")
            return False

    def _create_toc_field(self, position):
        """
        Crea un nuevo campo de tabla de contenido
        """
        # Crear párrafo para el campo TOC
        toc_paragraph = OxmlElement('w:p')

        # Configurar el campo TOC
        run = OxmlElement('w:r')
        fld_char = OxmlElement('w:fldChar')
        fld_char.set(qn('w:fldCharType'), 'begin')
        run.append(fld_char)
        toc_paragraph.append(run)

        run = OxmlElement('w:r')
        instr_text = OxmlElement('w:instrText')
        instr_text.set(qn('xml:space'), 'preserve')
        instr_text.text = 'TOC \\o "1-3" \\h \\z \\u'
        run.append(instr_text)
        toc_paragraph.append(run)

        run = OxmlElement('w:r')
        fld_char = OxmlElement('w:fldChar')
        fld_char.set(qn('w:fldCharType'), 'separate')
        run.append(fld_char)
        toc_paragraph.append(run)

        run = OxmlElement('w:r')
        fld_char = OxmlElement('w:fldChar')
        fld_char.set(qn('w:fldCharType'), 'end')
        run.append(fld_char)
        toc_paragraph.append(run)

        # Insertar en la posición especificada
        self.doc.element.body.insert(position, toc_paragraph)

    def update_table_of_contents(self):
        """
        Actualiza la tabla de contenido del documento Word
        """
        try:
            # Buscar y actualizar todos los campos TOC (Tabla de Contenido)
            for paragraph in self.doc.paragraphs:
                for run in paragraph.runs:
                    if run._element.xpath('.//w:fldChar[@w:fldCharType="begin"]'):
                        # Encontramos un campo, probablemente el TOC
                        print("✅ Campo de tabla de contenido encontrado, actualizando...")

                        # Forzar la actualización del campo
                        self._update_toc_field()
                        return True

            print("⚠️ No se encontró campo de tabla de contenido")
            return False

        except Exception as e:
            print(f"❌ Error al actualizar tabla de contenido: {e}")
            return False

    def _update_toc_field(self):
        """
        Actualiza el campo de tabla de contenido mediante XML
        """
        # Buscar el elemento del campo TOC
        for element in self.doc.element.body.iter():
            if element.tag.endswith('}p'):  # Párrafo
                for child in element.iter():
                    if child.tag.endswith('}fldChar') and child.get(qn('w:fldCharType')) == 'begin':
                        # Encontramos el inicio de un campo
                        field_element = element

                        # Agregar instrucción de actualización
                        for run in field_element.findall('.//w:r', {
                            'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                            for t in run.findall('.//w:t',
                                                 {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                                if 'TOC' in t.text:
                                    # Forzar actualización modificando el XML
                                    instr_text = qn('w:instrText')
                                    for instr in field_element.findall('.//w:instrText', {
                                        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                                        if 'TOC' in instr.text:
                                            instr.text = instr.text.replace('TOC', 'TOC \\u \\* MERGEFORMAT')
                                            print("✅ Campo TOC modificado para actualización")
                                            return True
        return False

    def insert_normative_text(self):
        """
        Inserta texto normativo DESPUÉS del título "2.1 MARCO NORMATIVO" SIN borrarlo
        """

        # Buscar título "2.1 MARCO NORMATIVO"
        for i, paragraph in enumerate(self.doc.paragraphs):
            text_upper = paragraph.text.upper()

            if any(pattern in text_upper for pattern in [
                "2.1 MARCO NORMATIVO",
                "MARCO NORMATIVO",
                "2.1MARCO NORMATIVO",
                "2.1 MARCO"
            ]):
                print(f"Encontrado MARCO NORMATIVO: '{paragraph.text}' en párrafo {i}")

                # Crear una lista con todo el contenido a insertar
                content_to_insert = []

                # 1. Párrafo introductorio
                content_to_insert.append({
                    'text': "\nPara la realización del siguiente estudio, se tiene en cuenta la siguiente normatividad:",
                    'type': 'normal'
                })

                # 2. Línea vacía
                content_to_insert.append({'text': "", 'type': 'empty'})

                # 3. Decreto
                content_to_insert.append({
                    'text': "Decreto 1076 de 2015 (MADS). Por medio del cual se expide el Decreto Único Reglamentario del Sector Ambiente y Desarrollo Sostenible.",
                    'type': 'decreto'
                })

                # 4. Línea vacía
                content_to_insert.append({'text': "", 'type': 'empty'})

                # 5. Artículos
                articles = [
                    "Artículo 2.2.3.3.9.3. Tratamiento convencional y criterios de calidad para consumo humano y doméstico. Los criterios de calidad admisibles para la destinación del recurso para consumo humano y doméstico son los que se relacionan a continuación, e indican que para su potabilización se requiere solamente tratamiento convencional.",
                    "Artículo 2.2.3.3.9.4. Desinfección y criterios de calidad para consumo humano y doméstico. Los criterios de calidad admisibles para la destinación del recurso para consumo humano y doméstico que para su potabilización se requiere sólo desinfección.",
                    "Artículo 2.2.3.3.9.5. Criterios de calidad para uso agrícola. Los criterios de calidad admisibles para la destinación del recurso para uso agrícola.",
                    "Artículo 2.2.3.3.9.6. Criterios de calidad para uso pecuario. Los criterios de calidad admisibles para la destinación del recurso para uso pecuario.",
                    "Artículo 2.2.3.3.9.10. Criterios de calidad para la preservación de flora y fauna. Los criterios de calidad admisibles para la preservación de flora y fauna."
                ]

                for article in articles:
                    content_to_insert.append({
                        'text': f"- {article}",
                        'type': 'article'
                    })
                    content_to_insert.append({'text': "", 'type': 'empty'})

                # Insertar contenido DESPUÉS del título (sin tocarlo)
                insertion_index = i + 1

                for j, content in enumerate(content_to_insert):
                    target_index = insertion_index + j

                    if target_index < len(self.doc.paragraphs):
                        next_para = self.doc.paragraphs[target_index]
                        new_para = next_para.insert_paragraph_before(content['text'])
                    else:
                        new_para = self.doc.add_paragraph(content['text'])

                    # Aplicar formato según el tipo
                    if content['type'] == 'normal':
                        new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        for run in new_para.runs:
                            run.font.name = self.font
                            run.font.size = Pt(11)
                            run.bold = False

                    elif content['type'] == 'decreto':
                        # Limpiar párrafo y recrear con formato mixto
                        for run in new_para.runs:
                            run._element.getparent().remove(run._element)

                        run_bold = new_para.add_run("Decreto 1076 de 2015 (MADS).")
                        run_bold.font.name = self.font
                        run_bold.font.size = Pt(11)
                        run_bold.bold = True

                        run_normal = new_para.add_run(
                            " Por medio del cual se expide el Decreto Único Reglamentario del Sector Ambiente y Desarrollo Sostenible.")
                        run_normal.font.name = self.font
                        run_normal.font.size = Pt(11)
                        run_normal.bold = False

                        new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    elif content['type'] == 'article':
                        # Limpiar párrafo y recrear con formato mixto
                        for run in new_para.runs:
                            run._element.getparent().remove(run._element)

                        text = content['text'][2:]  # Quitar "- "
                        parts = text.split(". ", 1)

                        if len(parts) > 1:
                            run_bold = new_para.add_run(f"- {parts[0]}.")
                            run_bold.font.name = self.font
                            run_bold.font.size = Pt(11)
                            run_bold.bold = True

                            run_normal = new_para.add_run(f" {parts[1]}")
                            run_normal.font.name = self.font
                            run_normal.font.size = Pt(11)
                            run_normal.bold = False

                        new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        new_para.paragraph_format.left_indent = Pt(18)
                        new_para.paragraph_format.space_before = Pt(3)
                        new_para.paragraph_format.space_after = Pt(3)

                    elif content['type'] == 'empty':
                        new_para.paragraph_format.space_after = Pt(6)

                print("✅ Texto normativo insertado DESPUÉS del título '2.1 MARCO NORMATIVO' (título conservado)")
                return True

        print("❌ No se encontró el título '2.1 MARCO NORMATIVO'")
        return False





    def write_methodology_section(self):

        methodology_found = False


        for i, paragraph in enumerate(self.doc.paragraphs):

            paragraph_text = paragraph.text.strip().upper()


            if any(pattern in paragraph_text for pattern in [

                "2. DESCRIPCION METODOLOGICA",
                "2. DESCRIPCIÓN METODOLÓGICA",
                "2.DESCRIPCION METODOLOGICA",
                "DESCRIPCION METODOLOGICA",
                "DESCRIPCIÓN METODOLÓGICA"

            ]) :

                print(f"ENCONTRADO TITULO METODOLOGIA")

                methodology_found = True

                target_para = True

                if i + 1  < len(self.doc.paragraphs):

                    next_para = self.doc.paragraphs[i + 1]
                    if not next_para.text.strip() or len(next_para.text.strip()) < 20:

                        target_para = next_para

                    else :

                        target_para = next_para.insert_paragraph_before("")
            else:


                target_para= self.doc.add_paragraph("")


            self.clear_paragraph(target_para)


            run = target_para.add_run("2.1 DESCRIPCIÓN DE LOS PUNTOS DE MONITOREO")
            run.font.name = self.font
            run.font.size = Pt(12)
            run.bold = True

            target_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            target_para.paragraph_format.space_before = Pt(12)
            target_para.paragraph_format.space_after = Pt(6)
            target_para.paragraph_format.left_indent = Pt(0)



            print("Titulo 2.1 insertado")
            break

        marco_normativo_found = False
        paragraphs_to_remove= []


        for paragraph in self.doc.paragraphs:

            paragraph_text = paragraph.text.strip().upper()


            if any(pattern in paragraph_text for pattern in [

                "MARCO NORMATIVO",
                "MARCO NORMATIVO:",
                "3. MARCO NORMATIVO",
                "2. MARCO NORMATIVO",
                "MARCO LEGAL"
            ]):


                print("TEXTO MARCO NORMATIVO ENCONTRADO")
                paragraphs_to_remove.append(paragraph)
                marco_normativo_found = True

        for paragraph in paragraphs_to_remove:

            self.remove_paragraph(paragraph)
            print("MARCO NORMATIVO ELIMINADO")

        if not marco_normativo_found:

            print("No se encontro marco normativo")

        return methodology_found


    def clear_paragraph(self, paragraph):

        for run in paragraph.runs[:]:

            run._element.getparent().remove(run._element)

    def remove_paragraph(self, paragraph):


        p_element = paragraph._element
        p_element.getparent().remove(p_element)

    def fill_monitoring_table_data_SIMPLE(self):
        """
        Versión CORREGIDA - Escribe datos en las filas correctas después de duplicar bloques
        """
        # Buscar la tabla
        table_found = None
        target_title = "Descripción del punto de monitoreo"

        for element in self.doc.element.body:
            if element.tag.endswith('p'):
                para_text = ""
                for run in element.findall('.//w:t',
                                           {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    if run.text:
                        para_text += run.text

                if target_title.lower() in para_text.lower():
                    next_element = element.getnext()
                    while next_element is not None:
                        if next_element.tag.endswith('tbl'):
                            for i, table in enumerate(self.doc.tables):
                                if table._element == next_element:
                                    table_found = table
                                    break
                            break
                        next_element = next_element.getnext()
                    break

        if table_found is None:
            print("❌ No se encontró la tabla")
            return False

        print(f"✅ Tabla encontrada con {len(table_found.rows)} filas")

        # Identificar el patrón de filas en la tabla
        # Las filas de datos suelen estar en posiciones específicas dentro de cada bloque
        sample_list = list(self.samples.items())

        # Primero, identifiquemos qué filas contienen datos (no encabezados)
        data_rows = []
        for i, row in enumerate(table_found.rows):
            row_text = ' '.join([cell.text.strip() for cell in row.cells]).lower()

            # Buscar filas que parecen contener datos (no títulos)
            if (any(keyword in row_text for keyword in ['xxx', 'placeholder', '...', 'ejemplo']) or
                    (len(row_text) > 0 and len(row_text) < 50 and
                     not any(header in row_text for header in
                             ['código', 'fecha', 'hora', 'identificación', 'coordenadas', 'fotografía']))):
                data_rows.append(i)

        print(f"Filas identificadas como datos: {data_rows}")

        # Si no encontramos filas de datos por patrones, usar el patrón esperado
        if not data_rows:
            print("⚠️ No se pudieron identificar filas de datos por patrones, usando patrón esperado")
            # Asumir que cada bloque de 4 filas tiene 1 fila de datos en la posición 2
            for i in range(len(sample_list)):
                data_row_index = 2 + (i * 4)
                if data_row_index < len(table_found.rows):
                    data_rows.append(data_row_index)

        # Escribir datos en las filas identificadas
        for i, (sample_id, sample_data) in enumerate(sample_list):
            if i >= len(data_rows):
                print(f"❌ No hay suficientes filas de datos para todas las muestras")
                break

            target_row = data_rows[i]
            print(f"Muestra {i + 1}: Escribiendo en fila {target_row + 1} (índice {target_row})")

            if target_row >= len(table_found.rows):
                print(f"❌ Fila {target_row + 1} no existe. Tabla solo tiene {len(table_found.rows)} filas")
                break

            row = table_found.rows[target_row]

            # Datos a escribir
            data_to_write = [
                sample_data.get('chemilab_code', ''),
                sample_data.get('sample_date', ''),
                sample_data.get('sample_time', ''),
                sample_data.get('sample_identification', ''),
                sample_data.get('coordinates_east', ''),
                sample_data.get('coordinates_north', ''),
                sample_data.get('photography', '')
            ]

            # Escribir en cada celda
            for col, data in enumerate(data_to_write):
                if col < len(row.cells):
                    try:
                        cell = row.cells[col]
                        current_text = cell.text.strip()

                        # Solo escribir si la celda está vacía o tiene placeholder
                        if (not current_text or
                                current_text.lower() in ['xxx', 'placeholder', '...', 'ejemplo'] or
                                len(current_text) < 3):

                            self.write_cell_safe(
                                cell,
                                str(data).upper() if data else "",
                                size=10,
                                bold=False,
                                align=WD_ALIGN_PARAGRAPH.CENTER
                            )
                            print(f"  ✓ Celda {col}: '{data}'")
                        else:
                            print(f"  ⚠️ Celda {col} ya tiene contenido: '{current_text}' - No se sobrescribe")

                    except Exception as e:
                        print(f"  ❌ Error en celda {col}: {e}")
                else:
                    print(f"  ⚠️ Celda {col} no existe")

            print(f"  ✅ Muestra {i + 1} procesada")

        print("🎉 PROCESO COMPLETADO")
        return True

    def write_cell_safe(self, cell, text, size=8, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=0):
        """
        Versión SEGURA de write_cell que NO borra contenido existente
        Solo limpia si la celda está realmente vacía
        """

        # Verificar si la celda tiene contenido importante
        current_text = cell.text.strip()

        # Si tiene contenido significativo, no tocar
        if (current_text and
                len(current_text) > 3 and
                current_text.lower() not in ["placeholder", "xxx", "...", "n/a"]):
            print(f"    🛡️ Protegiendo contenido: '{current_text}'")
            return False

        # Solo limpiar si es seguro
        self.clear_cell_safe(cell)

        # Agregar el nuevo contenido
        p = cell.add_paragraph()

        if isinstance(text, list):
            for text_part, is_bold in text:
                run = p.add_run(text_part)
                run.font.name = self.font
                run.font.size = Pt(size)
                run.bold = is_bold
        else:
            run = p.add_run(text)
            run.font.name = self.font
            run.font.size = Pt(size)
            run.bold = bold

        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1

        return True

    def clear_cell_safe(self, cell):
        """
        Versión SEGURA de clear_cell que preserva imágenes Y títulos importantes
        """

        # Guardar imágenes
        images = []
        important_content = []

        for paragraph in cell.paragraphs:
            para_text = paragraph.text.strip()

            # Preservar contenido importante (títulos, headers, etc.)
            if (para_text and
                    len(para_text) > 3 and
                    para_text.lower() not in ["placeholder", "xxx", "...", "n/a"]):
                # Es contenido importante, guardarlo
                important_content.append(paragraph)
                continue

            # Guardar imágenes
            for run in paragraph.runs:
                if run._element.xpath('.//w:drawing'):
                    images.extend(run._element.xpath('.//w:drawing'))

        # Solo eliminar párrafos que NO son importantes
        paragraphs_to_remove = []
        for p in cell.paragraphs:
            if p not in important_content:
                paragraphs_to_remove.append(p)

        # Eliminar solo párrafos no importantes
        for p in paragraphs_to_remove:
            p._element.getparent().remove(p._element)

        # Restaurar imágenes si había
        if images:
            if not cell.paragraphs:  # Si no quedan párrafos
                p = cell.add_paragraph()
            else:
                p = cell.paragraphs[0]

            for img in images:
                p._element.append(img)

    def duplicate_monitoring_blocks_SAFE(self):
        """
        Versión SEGURA que duplica sin romper el contenido existente
        """

        # Tu código de búsqueda de tabla (mantener igual)
        target_title = "Descripción del punto de monitoreo"
        table_found = None

        for element in self.doc.element.body:
            if element.tag.endswith('p'):
                para_text = ""
                for run in element.findall('.//w:t',
                                           {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    if run.text:
                        para_text += run.text

                if target_title.lower() in para_text.lower():
                    next_element = element.getnext()
                    while next_element is not None:
                        if next_element.tag.endswith('tbl'):
                            for i, table in enumerate(self.doc.tables):
                                if table._element == next_element:
                                    table_found = table
                                    break
                            break
                        next_element = next_element.getnext()
                    break

        if table_found is None:
            print(f"No se encontró tabla después del título '{target_title}'")
            return False

        original_rows = len(table_found.rows)
        template_block_size = 4

        print(f"Tabla encontrada con {original_rows} filas")

        if original_rows < template_block_size:
            print(f"Error: La tabla original tiene menos de {template_block_size} filas")
            return False

        # Guardar template del primer bloque
        template_rows = []
        for i in range(template_block_size):
            if i < len(table_found.rows):
                row_element = table_found.rows[i]._element
                template_rows.append(copy.deepcopy(row_element))

        # Calcular bloques adicionales
        samples_count = len(self.samples)
        additional_blocks_needed = samples_count - 1

        print(f"Necesitamos {additional_blocks_needed} bloques adicionales para {samples_count} muestras")

        # Agregar bloques adicionales
        for block_num in range(additional_blocks_needed):
            print(f"Creando bloque adicional {block_num + 2} de {samples_count}")

            for template_row in template_rows:
                new_row_element = copy.deepcopy(template_row)

                # LIMPIAR SOLO DATOS, NO TÍTULOS
                for cell in new_row_element.findall('.//w:tc',
                                                    {
                                                        'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    for para in cell.findall('.//w:p',
                                             {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                        para_text = ""
                        for text_elem in para.findall('.//w:t',
                                                      {
                                                          'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                            if text_elem.text:
                                para_text += text_elem.text

                        # SOLO limpiar si NO es un título importante
                        if (para_text and
                                not any(keyword in para_text.upper() for keyword in
                                        ['PLAN DE MUESTREO', 'DESCRIPCIÓN DEL PUNTO', 'CONDICIONES AMBIENTALES',
                                         'CÓDIGO', 'FECHA', 'HORA', 'IDENTIFICACIÓN', 'COORDENADAS', 'FOTOGRAFÍA',
                                         'ESTE', 'NORTE', 'MUESTRA']) and
                                len(para_text.strip()) < 10):  # Solo limpiar textos cortos

                            for text_elem in para.findall('.//w:t',
                                                          {
                                                              'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                                text_elem.text = ""

                # Agregar la fila a la tabla
                table_found._element.append(new_row_element)

        print(f"Bloques duplicados exitosamente. Tabla ahora tiene {len(table_found.rows)} filas")
        return True

    # 1. REEMPLAZA esta función:
    def setup_monitoring_table(self):
        print("=== CONFIGURANDO TABLA DE MONITOREO ===")

        # Usar las versiones SAFE
        if not self.duplicate_monitoring_blocks_SAFE():
            return False

        if not self.fill_monitoring_table_data_SIMPLE():
            return False

        return True

    def explore_tables(self, table, level = 0):

        prefix = "  " * level  # sangría según el nivel de anidación
        for i, row in enumerate(table.rows):
            row_text = [cell.text.strip() for cell in row.cells]
            print(f"{prefix}Fila {i}: {row_text}")

            for j, cell in enumerate(row.cells):
                if cell.tables:
                    print(f"{prefix}--> Subtabla encontrada en fila {i}, celda {j}")
                    for subtable in cell.tables:
                        self.explore_tables(subtable, level + 1)








    def save(self):

        print(self.samples)

        if self.doc:
            self.doc.save("templates/output.docx")
            print("Document saved")


    def open_word_document(self):
            """
            Abre el documento Word automáticamente (solo Windows)
            """
            try:
                output_path = "templates/output.docx"

                # Abrir directamente con os.startfile (funciona solo en Windows)
                os.startfile(output_path)
                print(f"✅ Documento abierto en Word: {output_path}")
                return True

            except Exception as e:
                print(f"❌ Error al abrir el documento: {e}")
                return False

    def save_and_open(self):
            """
            Guarda y abre el documento
            """
            self.save()  # Esto ya guarda el documento
            return self.open_word_document()