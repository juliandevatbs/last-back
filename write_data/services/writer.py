import json
from datetime import datetime

from docx import Document

from core.utils.data.date_literal import date_literal


class Writer():

    def __init__(self, config_path: str, template_path: str):
        self.config_path = config_path
        self.template_path = template_path
        self.json_config = None
        self.word_template = None

    def load_json_config(self):
        with open(self.config_path, 'r', encoding='utf-8') as json_file:
            self.json_config = json.load(json_file)

    def load_word_template(self):
        self.word_template = Document(self.template_path)

    def search_and_replace(self, label: str, to_write: str) -> int:
        replacements = 0
        print(f"🔍 Buscando label: {label}")

        try:
            print(" → Revisando párrafos...")
            for i, paragraph in enumerate(self.word_template.paragraphs):
                if i % 50 == 0:
                    print(f"   > Párrafo #{i}")
                if label in paragraph.text:
                    print(f"   ⚡ Reemplazo en párrafo -> {label}")
                    replacements += self._replace_in_paragraph(paragraph, label, to_write)

            print(" → Revisando tablas...")
            for t_index, table in enumerate(self.word_template.tables):
                print(f"   > Tabla #{t_index}")
                replacements += self._search_in_table(table, label, to_write)

            print(" → Revisando headers...")
            for s_index, section in enumerate(self.word_template.sections):
                print(f"   > Header #{s_index}")
                replaced = self._search_in_header(section.header, label, to_write)
                if replaced > 0:
                    print(f"   ⚡ Reemplazo en header -> {label}")
                    replacements += replaced

            print(f"✅ Fin de reemplazo para {label} ({replacements} reemplazos)")
            return replacements

        except Exception as e:
            print(f"❌ Error en search_and_replace para {label}: {e}")
            raise

    def _search_in_header(self, header, search_text, replace_text):
        replacements = 0

        for paragraph in header.paragraphs:
            if search_text in paragraph.text:
                replacements += self._replace_in_paragraph(paragraph, search_text, replace_text)

        for table in header.tables:
            replacements += self._search_in_table(table, search_text, replace_text)

        return replacements

    def _replace_in_paragraph(self, parrafo, texto_buscar, texto_reemplazo):
        replacements = 0

        for run in parrafo.runs:
            if texto_buscar in run.text:
                times = run.text.count(texto_buscar)
                if type(texto_reemplazo) != str:
                    texto_reemplazo = str(texto_reemplazo)
                run.text = run.text.replace(texto_buscar, texto_reemplazo)
                replacements += times

        return replacements

    def _search_in_table(self, table, texto_buscar, texto_reemplazo):
        replacements = 0

        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if texto_buscar in paragraph.text:
                        replacements += self._replace_in_paragraph(paragraph, texto_buscar, texto_reemplazo)

                if cell.tables:
                    for table_child in cell.tables:
                        replacements += self._search_in_table(table_child, texto_buscar, texto_reemplazo)

        return replacements

    from datetime import datetime
    from core.utils.data import date_literal



    def map_data_to_variables(self, basic_data, samples_data, specific_data):
            mapped_data = {}

            fecha_monitoreo = None
            if basic_data:
                fecha_monitoreo = basic_data.get("FECHA_MONITOREO", "")
                mapped_data["XX_INFORME_NUMERO_XX"] = basic_data.get("INFORME_NUMERO", "")

            if fecha_monitoreo:
                mapped_data["XX_FECHA_MONITOREO_XX"] = fecha_monitoreo

                try:
                    if isinstance(fecha_monitoreo, str):
                        fecha_dt = datetime.strptime(fecha_monitoreo, "%Y-%m-%d")
                    else:
                        fecha_dt = fecha_monitoreo

                    mapped_data["XX_MES_LITERAL_XX"] = fecha_dt.strftime("%B").capitalize()

                    mapped_data["XX_FECHA_MONITOREO_LITERAL_XX"] = date_literal(fecha_dt)

                except Exception as ex:
                    print(f"Error procesando fecha_monitoreo: {ex}")

            # 🔹 3. Valores que no vienen del Excel (se pueden parametrizar o dejar por defecto)
            mapped_data["XX_FECHA_ELABORACION_XX"] = datetime.now().strftime("%Y-%m-%d")
            mapped_data["XX_REVISADO_POR_XX"] = "Claudia Calderón"
            mapped_data["XX_ROL_REVISADOR_XX"] = "Profesional de proyectos"
            mapped_data["XX_AUTORIZADO_POR_XX"] = "Claudia Calderón"
            mapped_data["XX_AUTORIZADO_POR_ROL_XX"] = "Directora de Proyectos"

            return mapped_data

    def write_data(self, basic_data, samples_data, specific_data):
        mapped_data = self.map_data_to_variables(basic_data, samples_data, specific_data)


        for k, v in mapped_data.items():
            print(f"   {k}: {repr(v)}")

        for label, value in mapped_data.items():
            if value:
                self.search_and_replace(label, value)
            else:
                print(f"{label} está vacío, se omite")

    def save_document(self, output_path: str):
        self.word_template.save(output_path)