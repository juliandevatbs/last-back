from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from core.utils.data.hour_to_str import hour_to_str


def ph_table_writer(doc, font: str, size: int, bold: bool, ph_data: dict, target_title: str):
    """
    Writes pH data to the in-situ results table

    Args:
        doc: python-docx document
        ph_data: Dictionary with pH data
        font: Font name
        size: Font size
        bold: Whether text should be bold
        target_title: Title of the table to find

    Returns:
        bool: True if successful, False if error
    """
    try:
        print(f"Searching for table with title: '{target_title}'")

        ph_table = _find_ph_table(doc, target_title)
        if not ph_table:
            print(f"Table not found for title: {target_title}")
            return False

        print(f"Table found with {len(ph_table.rows)} rows")

        start_row_idx = _find_start_row(ph_table)
        if start_row_idx is None:
            return False

        col_indices = {
            'hour': 0,
            'ph': 1,
            'uncertainty': 2,
            'solids': -3,
            'flow': -1,
            'avg': 1,
            'avg_uncertainty': 2,
            'avg_reported': 4,
            'avg_flow': 6,
            'min_reported': 3,
            'min_reported_sedimentos': 4,
            'min_valor_reportado_caudales': 6,
            'max_reported': 3,
            'max_reported_sedimentos': 4,
            'max_valor_reportado_caudales': 6,
        }

        row_count = _write_measurements(ph_table, start_row_idx, ph_data, col_indices, font, size, bold)

        _write_averages(ph_table, start_row_idx, row_count, ph_data, col_indices, font, size, bold)

        _write_minimums(ph_table, start_row_idx, row_count, ph_data, col_indices, font, size, bold)

        _write_maximums(ph_table, start_row_idx, row_count, ph_data, col_indices, font, size, bold)

        print(f"Successfully written data to table: {target_title}")
        return True

    except Exception as e:
        print(f"Error writing to table {target_title}: {e}")
        import traceback
        traceback.print_exc()
        return False


def _find_start_row(ph_table):
    """
    Searches for the row containing XX_HORA_INSITU_XX
    Returns row index or None
    """
    for idx, row in enumerate(ph_table.rows):
        row_text = " ".join(cell.text.strip() for cell in row.cells)

        if "XX_HORA_INSITU_XX" in row_text or "XX_HORA_IN" in row_text:
            print(f"Found marker XX_HORA_INSITU_XX at row {idx}")
            return idx

    print("Marker XX_HORA_INSITU_XX not found in table")
    print("Showing all rows for debug:")
    for idx, row in enumerate(ph_table.rows):
        cells_content = [cell.text.strip()[:30] for cell in row.cells]
        print(f"  Row {idx}: {cells_content}")

    return None


def _write_measurements(ph_table, start_row_idx, ph_data, col_indices, font, size, bold):
    """
    Writes individual measurement data to the table
    Returns the number of rows written
    """
    row_count = 0

    for key, data in ph_data.items():
        if key == "_metadata":
            continue

        current_row_idx = start_row_idx + row_count

        if current_row_idx >= len(ph_table.rows):
            print(f"No more rows available. Wrote {row_count} records of {len(ph_data) - 1}")
            break

        current_row = ph_table.rows[current_row_idx]

        hour_str = hour_to_str(data.get('hour')) if data.get('hour') else ''
        ph_value = data.get('ph', '')
        uncertainty_str = _format_uncertainty(data.get('incertidumbre', ''))
        solids_value = data.get('solidos_sedimentables', '')
        flow_value = data.get('caudal', '')

        _write_cell_simple(current_row.cells[col_indices['hour']], hour_str, font, size, bold,
                           WD_ALIGN_PARAGRAPH.CENTER)
        _write_cell_simple(current_row.cells[col_indices['ph']], str(ph_value), font, size, bold,
                           WD_ALIGN_PARAGRAPH.CENTER)
        _write_cell_simple(current_row.cells[col_indices['uncertainty']], uncertainty_str, font, size, bold,
                           WD_ALIGN_PARAGRAPH.CENTER)
        _write_cell_simple(current_row.cells[col_indices['solids']], str(solids_value), font, size, bold,
                           WD_ALIGN_PARAGRAPH.CENTER)
        _write_cell_simple(current_row.cells[col_indices['flow']], str(flow_value), font, size, bold,
                           WD_ALIGN_PARAGRAPH.CENTER)

        print(f"  Row {current_row_idx}: Hour={hour_str}, pH={ph_value}, Uncertainty={uncertainty_str}, Solids={solids_value}, Flow={flow_value}")

        row_count += 1

    return row_count


def _write_averages(ph_table, start_row_idx, row_count, ph_data, col_indices, font, size, bold):
    """
    Writes the averages row
    """
    promedio_row = ph_table.rows[start_row_idx + row_count]
    metadata = ph_data.get("_metadata", {})

    media_value = metadata.get("media_valores", "")
    media_uncertainty = metadata.get("media_incertidumbre", "")
    media_flow = metadata.get("media_caudal", "")
    min_reported = metadata.get("valor_minimo_reportado", "")

    _write_cell_simple(promedio_row.cells[col_indices['avg']], f"{media_value:.3}", font, size, True,
                       WD_ALIGN_PARAGRAPH.CENTER)
    _write_cell_simple(promedio_row.cells[col_indices['avg_uncertainty']], media_uncertainty, font, size, bold,
                       WD_ALIGN_PARAGRAPH.CENTER)
    _write_cell_simple(promedio_row.cells[col_indices['avg_reported']], '<0,1', font, size, True,
                       WD_ALIGN_PARAGRAPH.CENTER)
    _write_cell_simple(promedio_row.cells[col_indices['avg_flow']], media_flow, font, size, True,
                       WD_ALIGN_PARAGRAPH.CENTER)
    _write_cell_simple(promedio_row.cells[col_indices['min_reported']], min_reported, font, size, bold,
                       WD_ALIGN_PARAGRAPH.CENTER)

    print(f"Averages written successfully")


def _write_minimums(ph_table, start_row_idx, row_count, ph_data, col_indices, font, size, bold):
    """
    Writes the minimums row
    """
    min_row = ph_table.rows[start_row_idx + row_count + 1]
    metadata = ph_data.get("_metadata", {})

    valor_minimo_reportado = metadata.get("valor_minimo_reportado", "")
    valor_minimo_reportado_sedimentos = metadata.get("min_valor_reportado_solidos", "")
    valor_minimo_reportado_caudales = metadata.get("min_valor_reportado_caudales", "")

    _write_cell_simple(min_row.cells[col_indices['min_reported']], valor_minimo_reportado, font, size, bold,
                       WD_ALIGN_PARAGRAPH.CENTER)
    _write_cell_simple(min_row.cells[col_indices['min_reported_sedimentos']], valor_minimo_reportado_sedimentos, font, size, bold,
                       WD_ALIGN_PARAGRAPH.CENTER)
    _write_cell_simple(min_row.cells[col_indices['min_valor_reportado_caudales']], valor_minimo_reportado_caudales, font,
                       size, bold, WD_ALIGN_PARAGRAPH.CENTER)

    print(f"Minimums written successfully")


def _write_maximums(ph_table, start_row_idx, row_count, ph_data, col_indices, font, size, bold):
    """
    Writes the maximums row
    """
    max_row = ph_table.rows[start_row_idx + row_count + 2]
    metadata = ph_data.get("_metadata", {})

    valor_maximo_reportado = metadata.get("valor_maximo_reportado", "")
    valor_maximo_reportado_sedimentos = metadata.get("max_valor_reportado_solidos", "")
    valor_maximo_reportado_caudales = metadata.get("max_valor_reportado_caudales", "")

    _write_cell_simple(max_row.cells[col_indices['max_reported']], valor_maximo_reportado, font, size, bold,
                       WD_ALIGN_PARAGRAPH.CENTER)
    _write_cell_simple(max_row.cells[col_indices['max_reported_sedimentos']], valor_maximo_reportado_sedimentos, font, size, bold,
                       WD_ALIGN_PARAGRAPH.CENTER)
    _write_cell_simple(max_row.cells[col_indices['max_valor_reportado_caudales']], valor_maximo_reportado_caudales, font,
                       size, bold, WD_ALIGN_PARAGRAPH.CENTER)

    print(f"Maximums written successfully")


def _format_uncertainty(uncertainty_value):
    """
    Formats uncertainty with the ± symbol
    """
    if not uncertainty_value or uncertainty_value == '':
        return ''

    if isinstance(uncertainty_value, (int, float)):
        return f"±{uncertainty_value:.4f}"

    uncertainty_str = str(uncertainty_value).strip()
    if not uncertainty_str.startswith('±'):
        return f"±{uncertainty_str}"

    return uncertainty_str


def _find_ph_table(doc, target_title):
    """
    Finds the in-situ results table containing XX_HORA_INSITU_XX
    """
    title_found = False
    tables_after_title = []

    for element in doc.element.body:
        if element.tag.endswith('p'):
            para_text = ""
            for run in element.findall('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                if run.text:
                    para_text += run.text

            if target_title.lower() in para_text.lower():
                print(f"Title found: '{para_text.strip()}'")
                title_found = True
                continue

        if title_found and element.tag.endswith('tbl'):
            for table in doc.tables:
                if table._element == element:
                    tables_after_title.append(table)
                    break

    print(f"Found {len(tables_after_title)} tables after title")

    for idx, table in enumerate(tables_after_title):
        print(f"\nChecking table #{idx + 1} ({len(table.rows)} rows, {len(table.columns)} columns)")

        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                if "XX_HORA_INSITU_XX" in cell.text or "XX_HORA_IN" in cell.text:
                    print(f"Correct table found! (Table #{idx + 1}, has marker)")
                    return table

        if table.rows:
            first_row_text = " | ".join([cell.text.strip()[:20] for cell in table.rows[0].cells])
            print(f"   First row: {first_row_text}")

    print("No table found with XX_HORA_INSITU_XX")
    return None


def _write_cell_simple(cell, text, font_name, font_size, bold, align):
    """
    Writes simple text to a cell with formatting
    """
    for paragraph in cell.paragraphs:
        paragraph.clear()

    if not cell.paragraphs:
        p = cell.add_paragraph()
    else:
        p = cell.paragraphs[0]

    run = p.add_run(str(text))
    run.font.name = font_name
    run.font.size = Pt(font_size)
    run.bold = bold

    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1