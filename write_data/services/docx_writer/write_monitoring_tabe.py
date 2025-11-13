from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import glob


def write_monitoring_table(doc, font: str, size: int, bold: bool, samples_data: dict, basic_data: dict,
                           table_config: dict):
    target_title = table_config["TITULO_BUSQUEDA"]
    keywords = table_config["PALABRAS_CLAVE_ENCABEZADO"]
    min_keywords = table_config["MIN_PALABRAS_CLAVE"]

    monitoring_tables = _find_monitoring_tables(doc, target_title, keywords, min_keywords)

    if not monitoring_tables:
        return False


    FIRST_DATA_ROW = table_config["FILA_PRIMER_DATO"]
    ROWS_PER_SAMPLE = table_config["FILAS_POR_MUESTRA"]
    cols = table_config["COLUMNAS"]
    image_folders = table_config["CARPETAS_IMAGENES"]

    sample_list = [(k, v) for k, v in samples_data.items() if k != 'OSI']

    available_rows_first_table = len(monitoring_tables[0].rows) - FIRST_DATA_ROW
    samples_in_first_table = available_rows_first_table // ROWS_PER_SAMPLE


    samples_in_second_table = 0
    if len(monitoring_tables) > 1:
        available_rows_second_table = len(monitoring_tables[1].rows) - FIRST_DATA_ROW
        samples_in_second_table = available_rows_second_table // ROWS_PER_SAMPLE

    for i, (sample_id, sample_data) in enumerate(sample_list):
        if i < samples_in_first_table:
            current_table = monitoring_tables[0]
            sample_index_in_table = i
            table_type = "PRIMERA"
        elif len(monitoring_tables) > 1 and i < samples_in_first_table + samples_in_second_table:
            current_table = monitoring_tables[1]
            sample_index_in_table = i - samples_in_first_table
            table_type = "SEGUNDA"
        else:
            break

        data_row_idx = FIRST_DATA_ROW + (sample_index_in_table * ROWS_PER_SAMPLE)
        desc_row_idx = data_row_idx + 1



        if sample_index_in_table == 0:
            prev_row_idx = data_row_idx - 1
            if prev_row_idx >= FIRST_DATA_ROW - 1:
                prev_row = current_table.rows[prev_row_idx]
                write_cell_safe(
                    prev_row.cells[0],
                    [(f"PLAN DE MUESTREO: {basic_data['client_data']['XX_PLAN_MUESTRO_AGUAS_XX']}", True)],
                    font, size, False, WD_ALIGN_PARAGRAPH.LEFT
                )

        if data_row_idx >= len(current_table.rows):
            break

        data_row = current_table.rows[data_row_idx]
        num_columns = len(data_row.cells)

        sample_date = f"{sample_data.get('sample_day', ''):02d}/{sample_data.get('sample_month', ''):02d}/{sample_data.get('sample_year', ''):02d}"

        write_cell_safe(
            data_row.cells[cols["CODIGO_MUESTRA"]],
            sample_data.get('chemilab_code', ''),
            font, size, False, WD_ALIGN_PARAGRAPH.CENTER
        )

        write_cell_safe(
            data_row.cells[cols["FECHA"]],
            sample_date,
            font, size, False, WD_ALIGN_PARAGRAPH.CENTER
        )

        write_cell_safe(
            data_row.cells[cols["HORA"]],
            sample_data.get('sampler_hour', ''),
            font, size, False, WD_ALIGN_PARAGRAPH.CENTER
        )

        write_cell_safe(
            data_row.cells[cols["NOMBRE"]],
            [(f"\n{sample_data.get('sample_identification', '')}\n", False)],
            font, size, False, WD_ALIGN_PARAGRAPH.CENTER
        )

        if cols["FOTOGRAFIA"] < num_columns:
            image_path = _get_sample_image(i, image_folders)
            if image_path and os.path.exists(image_path):
                try:
                    photo_cell = data_row.cells[cols["FOTOGRAFIA"]]
                    clear_cell_safe(photo_cell)

                    if not photo_cell.paragraphs:
                        photo_cell.add_paragraph()

                    paragraph = photo_cell.paragraphs[0]
                    run = paragraph.add_run()
                    run.add_picture(image_path, width=Inches(1.5))
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                except Exception as e:
                    print(f"Error al insertar imagen: {e}")

        if desc_row_idx < len(current_table.rows):
            desc_row = current_table.rows[desc_row_idx]
            description_text = sample_data.get('sample_description', '')

            write_cell_safe(
                desc_row.cells[cols["DESCRIPCION"]],
                [('Descripción del punto: ', True),
                 (f"{description_text[2:]}\n", False),
                 ("\nCondiciones ambientales: ", True),
                 (
                     f"{sample_data.get('sample_weather')} - Temperatura ambiente: {sample_data.get('sample_temperature')} - Humedad relativa: {sample_data.get('sample_humidity')} - Altitud: {sample_data.get('sample_altitude')}")
                 ],
                font, size, False, WD_ALIGN_PARAGRAPH.LEFT
            )

        print(f"  ✅ Muestra {sample_id} escrita exitosamente")

    return True


def _find_monitoring_tables(doc, target_title, keywords, min_keywords):
    monitoring_tables = []

    for element in doc.element.body:
        if element.tag.endswith('p'):
            para_text = ""
            for run in element.findall('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                if run.text:
                    para_text += run.text

            if target_title.lower() in para_text.lower():

                next_element = element.getnext()
                tables_found = 0

                while next_element is not None and tables_found < 3:
                    if next_element.tag.endswith('tbl'):
                        for table in doc.tables:
                            if table._element == next_element:
                                if _is_monitoring_table(table, keywords, min_keywords):
                                    monitoring_tables.append(table)
                                    tables_found += 1
                                break

                    next_element = next_element.getnext()

                break

    if not monitoring_tables:
        monitoring_tables = _find_tables_by_structure(doc, keywords, min_keywords)

    return monitoring_tables


def _is_monitoring_table(table, keywords, min_keywords):
    if len(table.rows) < 4:
        return False

    first_row_text = ""
    for cell in table.rows[0].cells:
        first_row_text += cell.text.upper() + " "

    found_keywords = sum(1 for keyword in keywords if keyword in first_row_text)

    return found_keywords >= min_keywords


def _find_tables_by_structure(doc, keywords, min_keywords):
    monitoring_tables = []

    for i, table in enumerate(doc.tables):
        if _is_monitoring_table(table, keywords, min_keywords):
            monitoring_tables.append(table)

    return monitoring_tables


def write_cell_safe(cell, text, font_name, font_size, bold, align, space_after=0):
    clear_cell_safe(cell)
    p = cell.add_paragraph()

    if isinstance(text, list):
        for item in text:
            if isinstance(item, (tuple, list)) and len(item) == 2:
                text_part, is_bold = item
                run = p.add_run(text_part)
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run.bold = is_bold
            else:
                run = p.add_run(str(item))
                run.font.name = font_name
                run.font.size = Pt(font_size)
                run.bold = bold
    else:
        run = p.add_run(str(text))
        run.font.name = font_name
        run.font.size = Pt(font_size)
        run.bold = bold

    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1


def clear_cell_safe(cell):
    images = []

    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            if run._element.xpath('.//w:drawing'):
                images.extend(run._element.xpath('.//w:drawing'))

    paragraphs_to_remove = list(cell.paragraphs)
    for p in paragraphs_to_remove:
        p._element.getparent().remove(p._element)

    if images:
        p = cell.add_paragraph()
        for img in images:
            p._element.append(img)


def _get_sample_image(sample_index, image_folders):
    folder_name = image_folders.get(str(sample_index))
    if not folder_name:
        return None

    current_dir = os.path.dirname(os.path.abspath(__file__))
    while os.path.basename(current_dir) != "BackEnd" and current_dir != os.path.dirname(current_dir):
        current_dir = os.path.dirname(current_dir)

    base_path = os.path.join(current_dir, "assets", "images", folder_name)

    if not os.path.exists(base_path):
        return None

    image_extensions = ['*.jpg', '*.jpeg', '*.png', '*.JPG', '*.JPEG', '*.PNG']

    for ext in image_extensions:
        images = glob.glob(os.path.join(base_path, ext))
        if images:
            images.sort()
            return images[0]

    return None