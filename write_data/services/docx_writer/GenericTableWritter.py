from docx.shared import Pt


class GenericTableWriter:
    """
    Escritor genérico de tablas basado en configuración JSON
    """

    def __init__(self, doc, config: dict, data: dict):
        self.doc = doc
        self.config = config
        self.data = data
        self.font = config.get('FUENTE', 'Century Gothic')
        self.size = config.get('TAMANO_FUENTE', 10)
        self.bold = config.get('NEGRITA', False)

    def find_table(self, table_config: dict):
        """
        Encuentra una tabla según el método especificado en el JSON
        """
        metodo = table_config.get('METODO_BUSQUEDA', 'por_indice')

        if metodo == 'por_indice':
            indice = table_config.get('INDICE_TABLA', 0)
            return self.doc.tables[indice] if indice < len(self.doc.tables) else None

        elif metodo == 'por_contenido':
            palabras_clave = table_config.get('PALABRAS_CLAVE_ENCABEZADO', [])
            min_palabras = table_config.get('MIN_PALABRAS_CLAVE', 1)
            return self._find_table_by_keywords(palabras_clave, min_palabras)

        elif metodo == 'por_titulo':
            titulo = table_config.get('TITULO_BUSQUEDA', '')
            return self._find_table_by_title(titulo)

        return None

    def _find_table_by_title(self, target_title: str):
        """
        Busca tabla después de encontrar un título en un párrafo
        """
        title_found = False

        for element in self.doc.element.body:
            if element.tag.endswith('p'):
                para_text = self._extract_text_from_element(element)

                if target_title.lower() in para_text.lower():
                    print(f"✓ Título encontrado: '{para_text.strip()}'")
                    title_found = True
                    continue

            if title_found and element.tag.endswith('tbl'):
                for table in self.doc.tables:
                    if table._element == element:
                        return table

        return None

    def _find_table_by_keywords(self, keywords: list, min_keywords: int):
        """
        Busca tabla que contenga ciertas palabras clave en su encabezado
        """
        for table in self.doc.tables:
            if not table.rows:
                continue

            first_row_text = " ".join(
                cell.text.upper() for cell in table.rows[0].cells
            )

            matches = sum(1 for keyword in keywords if keyword.upper() in first_row_text)

            if matches >= min_keywords:
                print(f"✓ Tabla encontrada con {matches} palabras clave coincidentes")
                return table

        return None

    def get_data_from_path(self, path: str):
        """
        Obtiene datos siguiendo una ruta como 'basic_data.fecha_monitoreo'
        """
        if path.startswith('computed.'):
            # Datos computados (ej: fecha actual)
            return self._compute_value(path.split('.')[1])

        parts = path.split('.')
        value = self.data

        for part in parts:
            if isinstance(value, dict):
                value = value.get(part)
            else:
                return None

        return value

    def _compute_value(self, key: str):
        """
        Calcula valores dinámicos
        """
        if key == 'fecha_actual':
            return datetime.now().strftime('%Y-%m-%d')
        return None

    def write_replacements(self, table, replacements: list):
        """
        Escribe reemplazos en una tabla según configuración
        """
        for replacement in replacements:
            placeholder = replacement.get('PLACEHOLDER')
            data_path = replacement.get('FUENTE_DATO')
            tipo = replacement.get('TIPO', 'texto')

            value = self.get_data_from_path(data_path)

            if value is None:
                print(f"⚠ Dato no encontrado para {data_path}")
                continue

            # Formatear según tipo
            if tipo == 'fecha' and isinstance(value, str):
                formato = replacement.get('FORMATO', '%Y-%m-%d')
                # Parsear y reformatear si es necesario
                value = value  # Simplificado

            # Buscar y reemplazar en toda la tabla
            self._replace_in_table(table, placeholder, str(value))

    def _replace_in_table(self, table, placeholder: str, value: str):
        """
        Reemplaza un placeholder en todas las celdas de una tabla
        """
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if placeholder in paragraph.text:
                        paragraph.text = paragraph.text.replace(placeholder, value)

                        for run in paragraph.runs:
                            run.font.name = self.font
                            run.font.size =Pt(self.size)
                            run.font.bold = self.bold

    def _extract_text_from_element(self, element):
        """
        Extrae texto de un elemento XML
        """
        para_text = ""
        for run in element.findall('.//w:t', {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
            if run.text:
                para_text += run.text
        return para_text