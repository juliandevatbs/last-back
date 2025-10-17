from datetime import datetime

from docx.shared import Pt, RGBColor

from core.utils.data.datetime_to_str import datetime_to_string


def write_first_page(doc, font: str, size: int, bold: bool, fecha_monitoreo: str):
    tables = doc.tables

    if len(tables) == 0:
        return False

    first_table = tables[0]


    # En el futuro hay que reemplazarlo por algo real informe_numero
    informe_numero = '127921'
    fecha_monitoreo = fecha_monitoreo
    fecha_emision =  today_date = datetime.now().strftime('%Y-%m-%d')

    table_replacements = {
        'XX_INFORME_NUMERO_XX': informe_numero,
        'XX_FECHA_MONITOREO_XX': fecha_monitoreo,
        'XX_FECHA_EMISION_INFORME_XX': fecha_emision
    }



    for row_idx, row in enumerate(first_table.rows):
        for col_idx, cell in enumerate(row.cells):

            for paragraph in cell.paragraphs:

                for placeholder, value in table_replacements.items():
                    if placeholder in paragraph.text:


                        if isinstance(value, datetime):
                            paragraph.text = paragraph.text.replace(placeholder, datetime_to_string(value))
                        else:
                            paragraph.text = paragraph.text.replace(placeholder, value)

                        for run in paragraph.runs:
                            run.font.name = font
                            run.font.size = Pt(size)
                            run.font.bold = bold





    cell = first_table.rows[1].cells[0]

    if len(cell.tables) == 0:
        return False

    subtable = cell.tables[0]

    replacements = {
        'XX_REVISADO_POR_XX': 'Andrea Carolina Barragán',
        'XX_ROL_REVISADOR_XX': 'Profesional de proyectos',
        'XX_AUTORIZADO_POR_XX': 'Claudia Calderón',
        'XX_AUTORIZADO_POR_ROL_XX': 'Directora de Proyectos'
    }

    for col_idx in range(3):
        subcell = subtable.rows[0].cells[col_idx]

        for paragraph in subcell.paragraphs:

            for placeholder, value in replacements.items():
                if placeholder in paragraph.text:

                    paragraph.text = paragraph.text.replace(placeholder, value)

                    for run in paragraph.runs:
                        run.font.name = font
                        run.font.size = Pt(size)
                        run.font.bold = bold


    print("\n✅ Todos los placeholders reemplazados")

    return True